"""
compiler.py
===========
ATS-Optimized Resume Document Compiler
----------------------------------------
Generates a single-column, machine-readable Word (.docx) resume from a
structured profile JSON, using python-docx with raw OpenXML overrides for
features the library doesn't expose via its Python API.

ATS Design Principles Applied:
  1. Single column layout — no tables, text boxes, or multi-column sections.
     ATS parsers read documents top-to-bottom, left-to-right. Columns,
     floating frames, and text boxes cause content to be scanned out of
     logical order, scrambling parsed output.
  2. Standard fonts — Calibri and Arial are universally recognized by
     all document parsers.
  3. Section borders — implemented via direct OpenXML injection (w:pBdr)
     since python-docx provides no native paragraph border API. These create
     clean visual dividers without using horizontal rule images or tables.
  4. Strict margins — 0.75-inch margins maximize readable content area while
     conforming to professional resume standards.
  5. Consistent heading hierarchy — H1 for name, H2-equivalent for section
     titles, normal paragraphs for body text.

Dependencies:
    pip install python-docx
"""

import logging
from copy import deepcopy
from io import BytesIO
from typing import Optional
from lxml import etree

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# ---------------------------------------------------------------------------
# Logging Setup
# ---------------------------------------------------------------------------
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s — %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
)
logger = logging.getLogger("compiler")

# ---------------------------------------------------------------------------
# Design Tokens
# ---------------------------------------------------------------------------

FONT_NAME = "Calibri"               # ATS-safe serif-less font
FONT_SIZE_NAME = Pt(20)             # Candidate name — prominent heading
FONT_SIZE_CONTACT = Pt(9)           # Contact line — compact
FONT_SIZE_SECTION = Pt(12)          # Section titles
FONT_SIZE_BODY = Pt(10)             # Body text
FONT_SIZE_BULLET = Pt(10)           # Bullet points

COLOR_NAME = RGBColor(0x1A, 0x1A, 0x2E)        # Near-black navy
COLOR_SECTION = RGBColor(0x16, 0x21, 0x3E)     # Dark navy for section headers
COLOR_COMPANY = RGBColor(0x0F, 0x3C, 0x78)     # Deep blue for company names
COLOR_BODY = RGBColor(0x33, 0x33, 0x33)        # Dark gray body text
COLOR_BORDER = "1A1A2E"                         # Section border hex (no #)

MARGIN_SIZE = Inches(0.75)  # 0.75-inch margins on all sides

# Bullet character — standard "•" maps correctly across all ATS systems.
BULLET_CHAR = "•"

# Separator for contact information line.
CONTACT_SEPARATOR = "  |  "


# ---------------------------------------------------------------------------
# OpenXML Helper Functions
# ---------------------------------------------------------------------------

def _set_paragraph_bottom_border(paragraph) -> None:
    """
    Inject a bottom border on a paragraph using raw OpenXML (w:pBdr/w:bottom).

    python-docx does not expose paragraph border properties through its public
    API, so we construct the XML elements manually and attach them to the
    paragraph's properties (pPr) element.

    Border specification:
      - style: "single" — a clean, solid 1pt line
      - sz:    "6"      — half-point units; 6 = 3pt visual weight
      - space: "1"      — 1pt gap between text and border
      - color: "#1A1A2E" — matches the section header color

    XML Structure inserted:
        <w:pPr>
          <w:pBdr>
            <w:bottom w:val="single" w:sz="6" w:space="1" w:color="1A1A2E"/>
          </w:pBdr>
        </w:pPr>

    Args:
        paragraph: A python-docx Paragraph object to apply the border to.
    """
    # Access or create the paragraph's properties XML element (pPr).
    pPr = paragraph._p.get_or_add_pPr()

    # Remove any existing border definition to avoid duplication on reruns.
    existing_pBdr = pPr.find(qn("w:pBdr"))
    if existing_pBdr is not None:
        pPr.remove(existing_pBdr)

    # Build the <w:pBdr> container element.
    pBdr = OxmlElement("w:pBdr")

    # Build the <w:bottom> border element with desired attributes.
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")       # Single solid line
    bottom.set(qn("w:sz"), "6")             # 3pt visual weight (half-points)
    bottom.set(qn("w:space"), "1")          # 1pt breathing room below text
    bottom.set(qn("w:color"), COLOR_BORDER) # Match section heading color

    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_cell_margins(cell, top=0, start=0, bottom=0, end=0) -> None:
    """
    Set internal cell margins for table cells (helper for potential future use).

    Args:
        cell:   A python-docx TableCell object.
        top/start/bottom/end: Margins in twips (1 inch = 1440 twips).
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def _apply_run_formatting(
    run,
    bold: bool = False,
    italic: bool = False,
    color: Optional[RGBColor] = None,
    size: Optional[Pt] = None,
    font_name: str = FONT_NAME,
) -> None:
    """
    Apply consistent font formatting to a python-docx Run object.

    Args:
        run:       A python-docx Run object (inline text element).
        bold:      Whether to apply bold weight.
        italic:    Whether to apply italic style.
        color:     RGBColor for the text color.
        size:      Font size as a Pt value.
        font_name: Font family name (defaults to FONT_NAME constant).
    """
    run.bold = bold
    run.italic = italic
    run.font.name = font_name
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = size


def _set_paragraph_spacing(
    paragraph,
    space_before: int = 0,
    space_after: int = 0,
    line_rule: str = "auto",
    line_spacing: int = 276,
) -> None:
    """
    Configure paragraph spacing via OpenXML w:spacing element.

    All values are in twips (1/20 of a point; 240 twips = 1 line = 12pt at
    default line height). This provides finer control than python-docx's
    paragraph_format.space_before/after attributes which truncate to integer pt.

    Args:
        paragraph:    python-docx Paragraph object.
        space_before: Space before paragraph in twips.
        space_after:  Space after paragraph in twips.
        line_rule:    Line spacing rule ("auto" | "exact" | "atLeast").
        line_spacing: Line spacing value in twips (240 = single, 480 = double).
    """
    pPr = paragraph._p.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:before"), str(space_before))
    spacing.set(qn("w:after"), str(space_after))
    spacing.set(qn("w:lineRule"), line_rule)
    spacing.set(qn("w:line"), str(line_spacing))


# ---------------------------------------------------------------------------
# Document Construction Functions
# ---------------------------------------------------------------------------

def _add_name_block(doc: Document, profile: dict) -> None:
    """
    Add the candidate's name as the document's primary heading.

    The name is centered, styled in a large dark-navy font, and followed by
    a single contact information line combining all contact fields.

    Args:
        doc:     The Document object being built.
        profile: The profile dict containing "name" and "contact" keys.
    """
    name = profile.get("name", "Candidate Name")
    contact = profile.get("contact", {})

    # -- Name paragraph ---------------------------------------------------
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(name_para, space_before=0, space_after=40)

    name_run = name_para.add_run(name)
    _apply_run_formatting(
        name_run, bold=True, color=COLOR_NAME, size=FONT_SIZE_NAME
    )

    # -- Professional Title line (Job-specific headline) -----------------
    title_text = profile.get("professional_title")
    if title_text:
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_spacing(title_para, space_before=20, space_after=40)
        title_run = title_para.add_run(title_text)
        _apply_run_formatting(
            title_run, bold=True, italic=True, color=COLOR_COMPANY, size=Pt(11)
        )

    # -- Contact line -----------------------------------------------------
    contact_parts = []
    if contact.get("email"):
        contact_parts.append(contact["email"])
    if contact.get("phone"):
        contact_parts.append(contact["phone"])
    if contact.get("linkedin"):
        contact_parts.append(contact["linkedin"])
    if contact.get("github"):
        contact_parts.append(contact["github"])
    if contact.get("location"):
        contact_parts.append(contact["location"])

    if contact_parts:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_spacing(contact_para, space_before=0, space_after=80)

        contact_run = contact_para.add_run(CONTACT_SEPARATOR.join(contact_parts))
        _apply_run_formatting(
            contact_run, color=COLOR_BODY, size=FONT_SIZE_CONTACT
        )


def _add_section_title(doc: Document, title: str) -> None:
    """
    Add a section title paragraph with an OpenXML bottom border.

    The border is injected via _set_paragraph_bottom_border() since python-docx
    has no native API for this. The section title text is all-caps, bold, and
    in the section header color.

    Args:
        doc:   The Document object being built.
        title: The section name (e.g., "Professional Experience").
    """
    para = doc.add_paragraph()
    _set_paragraph_spacing(para, space_before=160, space_after=60)

    run = para.add_run(title.upper())
    _apply_run_formatting(
        run, bold=True, color=COLOR_SECTION, size=FONT_SIZE_SECTION
    )

    # Inject the OpenXML bottom border — this is the key ATS-safe divider.
    _set_paragraph_bottom_border(para)


def _add_summary_section(doc: Document, profile: dict) -> None:
    """
    Add the Professional Summary section.

    Args:
        doc:     The Document object.
        profile: Profile dict with a "summary" string key.
    """
    summary_text = profile.get("summary", "")
    if not summary_text:
        return

    _add_section_title(doc, "Professional Summary")

    para = doc.add_paragraph()
    _set_paragraph_spacing(para, space_before=40, space_after=40)
    run = para.add_run(summary_text)
    _apply_run_formatting(run, color=COLOR_BODY, size=FONT_SIZE_BODY)


def _add_skills_section(doc: Document, profile: dict) -> None:
    """
    Add the Technical Skills section, grouped by category.

    Skills are rendered as "Category: skill1, skill2, skill3" lines to
    maintain ATS parseability while grouping related technologies.

    Args:
        doc:     The Document object.
        profile: Profile dict with "technical_skills" dict key.
    """
    skills = profile.get("technical_skills", {})
    if not skills:
        return

    _add_section_title(doc, "Technical Skills")

    for category, skill_list in skills.items():
        if not skill_list:
            continue

        para = doc.add_paragraph()
        _set_paragraph_spacing(para, space_before=20, space_after=20)

        # Category label in bold, skills in normal weight.
        cat_run = para.add_run(f"{category}: ")
        _apply_run_formatting(cat_run, bold=True, color=COLOR_COMPANY, size=FONT_SIZE_BODY)

        skills_str = ", ".join(skill_list) if isinstance(skill_list, list) else str(skill_list)
        skills_run = para.add_run(skills_str)
        _apply_run_formatting(skills_run, color=COLOR_BODY, size=FONT_SIZE_BODY)


def _add_experience_section(doc: Document, profile: dict) -> None:
    """
    Add the Professional Experience section.

    For each role, renders:
      - Job title (bold) + Company name (bold, colored) on one line.
      - Date range + location on a second line (italic).
      - Bullet points for each achievement.

    Args:
        doc:     The Document object.
        profile: Profile dict with "professional_experience" list key.
    """
    experiences = profile.get("professional_experience", [])
    if not experiences:
        return

    _add_section_title(doc, "Professional Experience")

    for i, exp in enumerate(experiences):
        # -- Role title + company line ------------------------------------
        title_para = doc.add_paragraph()
        _set_paragraph_spacing(
            title_para,
            space_before=100 if i > 0 else 60,
            space_after=20,
        )

        title_run = title_para.add_run(exp.get("title", "Role"))
        _apply_run_formatting(
            title_run, bold=True, color=COLOR_NAME, size=FONT_SIZE_BODY
        )

        separator_run = title_para.add_run("  —  ")
        _apply_run_formatting(separator_run, color=COLOR_BODY, size=FONT_SIZE_BODY)

        company_run = title_para.add_run(exp.get("company", "Company"))
        _apply_run_formatting(
            company_run, bold=True, color=COLOR_COMPANY, size=FONT_SIZE_BODY
        )

        # -- Date + location line -----------------------------------------
        meta_para = doc.add_paragraph()
        _set_paragraph_spacing(meta_para, space_before=0, space_after=40)

        dates = exp.get("dates", "")
        location = exp.get("location", "")
        meta_text = f"{dates}"
        if location:
            meta_text += f"  |  {location}"

        meta_run = meta_para.add_run(meta_text)
        _apply_run_formatting(
            meta_run, italic=True, color=COLOR_BODY, size=Pt(9)
        )

        # -- Achievement bullets ------------------------------------------
        for bullet in exp.get("bullets", []):
            bullet_para = doc.add_paragraph()
            _set_paragraph_spacing(bullet_para, space_before=20, space_after=20)
            bullet_para.paragraph_format.left_indent = Inches(0.2)

            bullet_run = bullet_para.add_run(f"{BULLET_CHAR}  {bullet}")
            _apply_run_formatting(bullet_run, color=COLOR_BODY, size=FONT_SIZE_BULLET)


def _add_education_section(doc: Document, profile: dict) -> None:
    """
    Add the Education section.

    Args:
        doc:     The Document object.
        profile: Profile dict with "education" list key.
    """
    education = profile.get("education", [])
    if not education:
        return

    _add_section_title(doc, "Education")

    for edu in education:
        # -- Degree + institution line ------------------------------------
        degree_para = doc.add_paragraph()
        _set_paragraph_spacing(degree_para, space_before=60, space_after=20)

        degree_run = degree_para.add_run(
            f"{edu.get('degree', '')} in {edu.get('field', '')}"
        )
        _apply_run_formatting(degree_run, bold=True, color=COLOR_NAME, size=FONT_SIZE_BODY)

        sep_run = degree_para.add_run("  —  ")
        _apply_run_formatting(sep_run, color=COLOR_BODY, size=FONT_SIZE_BODY)

        inst_run = degree_para.add_run(edu.get("institution", ""))
        _apply_run_formatting(inst_run, bold=True, color=COLOR_COMPANY, size=FONT_SIZE_BODY)

        # -- Graduation year + GPA (if provided) -------------------------
        meta_parts = []
        if edu.get("graduation_year"):
            meta_parts.append(str(edu["graduation_year"]))
        if edu.get("gpa"):
            meta_parts.append(f"GPA: {edu['gpa']}")
        if edu.get("honors"):
            meta_parts.append(edu["honors"])

        if meta_parts:
            meta_para = doc.add_paragraph()
            _set_paragraph_spacing(meta_para, space_before=0, space_after=20)
            meta_run = meta_para.add_run("  |  ".join(meta_parts))
            _apply_run_formatting(meta_run, italic=True, color=COLOR_BODY, size=Pt(9))

        # -- Relevant coursework / activities (optional) -----------------
        if edu.get("relevant_coursework"):
            cw_para = doc.add_paragraph()
            _set_paragraph_spacing(cw_para, space_before=10, space_after=10)
            cw_label = cw_para.add_run("Relevant Coursework: ")
            _apply_run_formatting(cw_label, bold=True, color=COLOR_BODY, size=Pt(9))
            cw_run = cw_para.add_run(", ".join(edu["relevant_coursework"]))
            _apply_run_formatting(cw_run, color=COLOR_BODY, size=Pt(9))


def _add_certifications_section(doc: Document, profile: dict) -> None:
    """
    Add an optional Certifications section.

    Args:
        doc:     The Document object.
        profile: Profile dict with optional "certifications" list key.
    """
    certs = profile.get("certifications", [])
    if not certs:
        return

    _add_section_title(doc, "Certifications")

    for cert in certs:
        cert_para = doc.add_paragraph()
        _set_paragraph_spacing(cert_para, space_before=40, space_after=20)
        cert_para.paragraph_format.left_indent = Inches(0.2)

        cert_run = cert_para.add_run(f"{BULLET_CHAR}  {cert.get('name', '')}")
        _apply_run_formatting(cert_run, bold=True, color=COLOR_BODY, size=FONT_SIZE_BODY)

        issuer = cert.get("issuer", "")
        year = cert.get("year", "")
        if issuer or year:
            meta = f"  ({issuer}{', ' + str(year) if year else ''})"
            meta_run = cert_para.add_run(meta)
            _apply_run_formatting(meta_run, italic=True, color=COLOR_BODY, size=Pt(9))


def _add_projects_section(doc: Document, profile: dict) -> None:
    """
    Add an optional Projects section.

    Args:
        doc:     The Document object.
        profile: Profile dict with optional "projects" list key.
    """
    projects = profile.get("projects", [])
    if not projects:
        return

    _add_section_title(doc, "Projects")

    for proj in projects:
        proj_para = doc.add_paragraph()
        _set_paragraph_spacing(proj_para, space_before=60, space_after=20)

        proj_name_run = proj_para.add_run(proj.get("name", "Project"))
        _apply_run_formatting(
            proj_name_run, bold=True, color=COLOR_NAME, size=FONT_SIZE_BODY
        )

        link = proj.get("live_link")
        if link:
            link_run = proj_para.add_run(f"  ({link})")
            _apply_run_formatting(link_run, italic=True, color=COLOR_COMPANY, size=Pt(9))

        tech = proj.get("technologies", "")
        if tech:
            tech_run = proj_para.add_run(f"  [{tech}]")
            _apply_run_formatting(tech_run, italic=True, color=COLOR_BODY, size=Pt(9))

        for bullet in proj.get("bullets", []):
            bp = doc.add_paragraph()
            _set_paragraph_spacing(bp, space_before=15, space_after=15)
            bp.paragraph_format.left_indent = Inches(0.2)
            br = bp.add_run(f"{BULLET_CHAR}  {bullet}")
            _apply_run_formatting(br, color=COLOR_BODY, size=FONT_SIZE_BULLET)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def compile_resume(profile: dict, output_path: str) -> str:
    """
    Compile a complete ATS-optimized resume document from a profile dict.

    This is the primary entry point called by core_orchestrator.py.

    The document is built in a strict sequential, single-column layout
    without any tables, columns, text boxes, or images that could confuse
    ATS parsing engines. OpenXML is injected directly for features
    (like paragraph borders) not exposed by python-docx's public API.

    Document Section Order:
      1. Name & Contact Block
      2. Professional Summary
      3. Technical Skills
      4. Professional Experience
      5. Education
      6. Certifications (if present)
      7. Projects (if present)

    Args:
        profile:     The (possibly tailored) profile dict. Expected to match
                     the schema defined in base_profile.json.
        output_path: Absolute file path where the .docx will be saved.

    Returns:
        The output_path string on success (for chaining in orchestrator).

    Raises:
        IOError: If the document cannot be saved to output_path.
        KeyError: If critical required profile keys are missing.
    """
    logger.info("Compiling resume to: %s", output_path)

    doc = Document()

    # -----------------------------------------------------------------------
    # Configure page margins
    # -----------------------------------------------------------------------
    # Word documents have one or more Section objects that control page layout.
    # The default document has exactly one section.
    section = doc.sections[0]
    section.top_margin = MARGIN_SIZE
    section.bottom_margin = MARGIN_SIZE
    section.left_margin = MARGIN_SIZE
    section.right_margin = MARGIN_SIZE

    # -----------------------------------------------------------------------
    # Remove default paragraph spacing from the "Normal" style to give us
    # full control via our own _set_paragraph_spacing calls.
    # -----------------------------------------------------------------------
    normal_style = doc.styles["Normal"]
    normal_style.font.name = FONT_NAME
    normal_style.font.size = FONT_SIZE_BODY
    normal_style.paragraph_format.space_before = Pt(0)
    normal_style.paragraph_format.space_after = Pt(0)

    # -----------------------------------------------------------------------
    # Build each resume section in order
    # -----------------------------------------------------------------------
    try:
        _add_name_block(doc, profile)
        _add_summary_section(doc, profile)
        _add_skills_section(doc, profile)
        _add_experience_section(doc, profile)
        _add_education_section(doc, profile)
        _add_certifications_section(doc, profile)
        _add_projects_section(doc, profile)
    except Exception as exc:
        logger.error("Error building document sections: %s", exc, exc_info=True)
        raise

    # -----------------------------------------------------------------------
    # Save DOCX and convert to PDF
    # -----------------------------------------------------------------------
    try:
        doc.save(output_path)
        logger.info("Resume DOCX saved successfully: %s", output_path)

        # Generate matching native PDF version
        pdf_path = output_path.replace(".docx", ".pdf")
        try:
            from pdf_compiler import compile_pdf_resume
            compile_pdf_resume(profile, pdf_path)
            logger.info("Native Resume PDF generated successfully: %s", pdf_path)
        except Exception as pe:
            logger.warning("Native PDF compile error: %s", pe)

    except IOError as exc:
        logger.error("Failed to save document to '%s': %s", output_path, exc)
        raise

    return output_path


def compile_resume_to_bytes(profile: dict) -> bytes:
    """
    Compile a resume and return it as raw bytes (for email attachment or API upload).

    Useful when you need the document content without writing to disk first.

    Args:
        profile: The profile dict.

    Returns:
        DOCX file content as bytes.
    """
    doc = Document()

    section = doc.sections[0]
    section.top_margin = MARGIN_SIZE
    section.bottom_margin = MARGIN_SIZE
    section.left_margin = MARGIN_SIZE
    section.right_margin = MARGIN_SIZE

    normal_style = doc.styles["Normal"]
    normal_style.font.name = FONT_NAME
    normal_style.font.size = FONT_SIZE_BODY
    normal_style.paragraph_format.space_before = Pt(0)
    normal_style.paragraph_format.space_after = Pt(0)

    _add_name_block(doc, profile)
    _add_summary_section(doc, profile)
    _add_skills_section(doc, profile)
    _add_experience_section(doc, profile)
    _add_education_section(doc, profile)
    _add_certifications_section(doc, profile)
    _add_projects_section(doc, profile)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
